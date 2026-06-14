# -*- coding: utf-8 -*-
"""Generate integrated 10-min presentation Word document for EMF Quantum Focus."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "EMF_Quantum_Focus_Presentation_Integrated.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("EMF Quantum Focus\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("스마트폰 센서 기반 학습 집중 환경 스캐너\n10분 발표 통합본")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run("\n버전 v1.0.0 · Flutter Android APK\n발표 시간: 약 10~11분 · 슬라이드 13장 권장")
    m.font.size = Pt(11)
    m.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Malgun Gothic"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_defaults(doc)

    # 목차
    add_title_page(doc)
    add_heading(doc, "목차", 1)
    toc = [
        "1. 표지 · 발표 개요",
        "2. 문제 인식 (Why)",
        "3. 독창성 (Differentiation)",
        "4. 핵심 기능 — 홈·집중",
        "5. 핵심 기능 — AR·맵·구역",
        "6. 동작 원리 — 센서·신호",
        "7. 동작 원리 — Focus Score·알림",
        "8. 동작 원리 — AR·구역 스캔",
        "9. 기술 스택",
        "10. 신뢰성 근거",
        "11. 한계점",
        "12. 보완점",
        "13. 확장 가능성 · 마무리",
        "부록: 발표 시간 배분 · PPT AI 프롬프트",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()

    # 1
    add_heading(doc, "슬라이드 1 — 표지 · 발표 개요", 1)
    add_bullets(doc, [
        "제목: EMF Quantum Focus",
        "부제: 스마트폰 센서 기반 학습 집중 환경 스캐너",
        "한 줄: 자기장·움직임을 재서 「지금 이 자리」의 집중 환경을 숫자와 맵으로 보여 주는 Android 앱",
        "버전: v1.0.0 (Flutter Release APK, 약 46.5MB)",
    ])
    add_quote(doc, "발표 한 줄: 「전문 계측기를 대체하지 않고, 일상 학습에서 환경을 의식하게 만드는 스마트폰 네이티브 집중 도구」")
    doc.add_page_break()

    # 2
    add_heading(doc, "슬라이드 2 — 문제 인식 (Why)", 1)
    add_bullets(doc, [
        "공부 방해 요인: 소음, 스마트폰 사용, 주변 전자기기 등이 복합적으로 작용",
        "기존 집중 앱: 타이머·앱 차단 중심 → 「지금 이 책상/자리」 환경은 측정하지 않음",
        "전문 EMF 계측기: 비용·휴대성·즉시 피드백 측면에서 일상 학습에 부적합",
        "목표: 스마트폰만으로 환경을 상대적으로 비교하고, 집중 습관을 돕는 도구",
    ])
    doc.add_page_break()

    # 3
    add_heading(doc, "슬라이드 3 — 독창성 (Differentiation)", 1)
    add_bullets(doc, [
        "EMF + 집중 UX 결합: 자기장 데이터 → Focus Score(0~100) → 집중 세션·알림 연동",
        "온디바이스 우선: 원시 센서는 기기 내 처리, 클라우드는 선택적 Gemini 요약만",
        "공간 시각화 3모드: 홈(점수·세션) / AR(카메라+오버레이) / 미니맵·구역(위치 탐색)",
        "8자 보정 + 상대 비교: 기기·환경 편차 보정, 같은 자세·같은 기기에서 전후 비교",
        "민감도 프리셋: 일반(≤60) / 도서관(≤63) / 민감(≤70) — 사용자 선택형 알림",
    ])
    add_quote(doc, "「집중 타이머 앱이 아니라, 내 자리의 환경을 재는 집중 코치」")
    doc.add_page_break()

    # 4
    add_heading(doc, "슬라이드 4 — 핵심 기능: 몰입 대시보드 (홈)", 1)
    add_table(doc, ["기능", "설명"], [
        ["Focus Score 게이지", "실시간 0~100 + 이중 링(세션 25분 목표 / 데이터 축적)"],
        ["집중 시작·종료", "세션 타이머, 하단 FocusSessionStrip, 통계 반영"],
        ["집중 알림 모드", "임계 이하 시 다이얼로그 + 맥락형 조치 안내"],
        ["오늘의 집중", "오늘/전체 분, 완료 세션 (SharedPreferences)"],
        ["측정 ON/OFF", "AppBar 센서 토글"],
    ])
    add_bullets(doc, ["UX: 다른 탭 이동 시에도 집중 세션·알림 유지"])
    doc.add_page_break()

    # 5
    add_heading(doc, "슬라이드 5 — 핵심 기능: AR·미니맵·구역", 1)
    add_heading(doc, "AR EMF 스캔", 2)
    add_bullets(doc, [
        "후면 카메라 실시간 영상 + EMF 오버레이(색·곡선·화살표·파티클)",
        "마그네토미터 dx, dy, dz → 방향 시각화",
    ])
    add_heading(doc, "EMF 미니맵", 2)
    add_bullets(doc, ["책상 평면 탑다운 맵", "이동하며 EMF 피크·완성도 누적"])
    add_heading(doc, "집중 구역 찾기", 2)
    add_bullets(doc, [
        "주변 스캔 후 Focus Score 최고 위치 표시",
        "범위: 수평 6m(±3m), 높이 ±1m (10m 요청 대비 센서·적분 한계로 축소)",
    ])
    doc.add_page_break()

    # 6
    add_heading(doc, "슬라이드 6 — 동작 원리: 센서·신호", 1)
    add_heading(doc, "사용 센서", 2)
    add_table(doc, ["센서", "역할"], [
        ["3축 마그네토미터", "EMF 핵심 (~100Hz 목표)"],
        ["선형 가속도", "책상/주변 위치 추정 (적분, SLAM 전)"],
        ["후면 카메라", "AR 배경 영상만 (영상 EMF 분석 없음)"],
    ])
    add_heading(doc, "신호 처리 파이프라인", 2)
    add_bullets(doc, [
        "8자 보정 오프셋 → 롤링 평균(DC 제거) → 축별 HPF",
        "크기(µT) + 이동평균(N=10) + 분산",
        "128점 FFT → 50/60Hz 전력선 대역 비율",
        "ICNIRP 스타일 UX 위험도 밴드",
    ])
    add_quote(doc, "중요: 폰 현재 위치에서만 자기장 측정 (원격 탐지 아님)")
    doc.add_page_break()

    # 7
    add_heading(doc, "슬라이드 7 — 동작 원리: Focus Score·알림", 1)
    p = doc.add_paragraph()
    r = p.add_run("Focus Score = 100 − EMF평균×0.98 − 분산×0.17 − 맵미완성×26 − 전력선%×0.26\n(집중 세션: 손 움직임 패널티 추가)")
    r.font.name = "Consolas"
    add_heading(doc, "알림 로직", 2)
    add_bullets(doc, [
        "1초 주기 점수 검사",
        "임계 이하 → 알림 (armed 해제)",
        "닫은 뒤 기준+4점 회복 후 재하락 시에만 재알림 (히스테리시스)",
    ])
    add_quote(doc, "의미: 절대 건강/성적 보장이 아닌, 같은 환경에서의 상대 비교·습관 형성용")
    doc.add_page_break()

    # 8
    add_heading(doc, "슬라이드 8 — 동작 원리: AR·구역 스캔", 1)
    add_bullets(doc, [
        "AR: CameraPreview + CustomPaint 2D 오버레이 (ARCore/ARKit 미사용, SLAM AR 아님)",
        "구역: 스캔 시작점=원점, 25cm 격자, 격자당 3샘플 이상 → 최고 Score 위치·방위·거리",
    ])
    doc.add_page_break()

    # 9
    add_heading(doc, "슬라이드 9 — 기술 스택", 1)
    add_table(doc, ["구분", "기술"], [
        ["앱", "Flutter / Dart (Material 3)"],
        ["상태", "provider — EnvironmentProvider"],
        ["센서", "sensors_plus (자기장, 가속도)"],
        ["카메라", "camera + permission_handler"],
        ["저장", "shared_preferences"],
        ["선택 AI", "http → Gemini API (요약 통계만)"],
        ["신호처리", "자체 FFT, HPF, 이동평균, ICNIRP UX"],
        ["배포", "Android APK Release"],
    ])
    add_quote(doc, "아키텍처: 센서 → Signal Processing → Provider → 4탭 UI")
    doc.add_page_break()

    # 10 reliability
    add_heading(doc, "슬라이드 10 — 신뢰성 근거", 1)
    add_quote(doc, "신뢰성 = 의학·학업 성과 보장이 아니라, 측정 일관성·투명한 한계 고지·반복 가능한 상대 비교")

    add_heading(doc, "1. 기술적 신뢰성", 2)
    add_table(doc, ["근거", "내용"], [
        ["8자 보정", "오프셋 저장·재적용 → 기준선 안정"],
        ["고정 파이프라인", "동일 규칙으로 숫자 산출"],
        ["Score 공식 명시", "코드에 가중치·임계값 고정"],
        ["히스테리시스 4점", "연속 오알림 완화"],
        ["격자 최소 3샘플", "구역 탐색 성급 판정 방지"],
        ["온디바이스 우선", "원시 스트림 기기 내, Gemini는 요약만"],
        ["한계 고지 UI", "계측기 아님·상대 비교용 명시"],
    ])

    add_heading(doc, "2. 학술·산업 근거 (간접)", 2)
    add_table(doc, ["주장", "근거", "주의"], [
        ["보정 필요", "마그네토미터 보정 문헌(NXP AN4246 등)", "측정 공학"],
        ["알림 설계", "인터럽션 비용 HCI(Mark CHI 등)", "Score 정확도 증명 아님"],
        ["단순 UI", "인지부하 이론(Sweller)", "정보 단순화"],
        ["EMF→집중", "WHO/ICNIRP 맥락", "인과 주장 금지"],
    ])

    add_heading(doc, "3. 검증 (현재)", 2)
    add_bullets(doc, [
        "flutter test 통과 (이동평균, 위젯, 구역 스캔)",
        "Release APK 빌드 성공",
        "사용자 연구·임상 시험: 없음 (상위 신뢰 근거는 다음 단계)",
    ])

    add_heading(doc, "4. 보장 범위", 2)
    add_bullets(doc, [
        "보장함: 동일 조건 반복 측정 일관성, 규칙 기반 알림, 한계 투명 고지",
        "보장 안 함: 절대 EMF 정확도, ICNIRP 준수, 학업 성과, AR SLAM 정밀도",
    ])
    doc.add_page_break()

    # 11
    add_heading(doc, "슬라이드 11 — 한계점", 1)
    add_bullets(doc, [
        "의료·계측기 아님 — ICNIRP 밴드는 UX 참고",
        "1점 측정 — 폰 위치만, 카메라 원격 EMF 측정 없음",
        "위치 추정 오차 — 가속도 적분, 장거리 드리프트",
        "AR는 SLAM AR 아님 — 3D 공간 고정 앵커 없음",
        "Score 가중치 — 벤치 튜닝 전, 기기별 편차",
        "Gemini — API 키 없으면 로컬 스텁",
    ])
    add_quote(doc, "한계를 숨기지 않는 것 자체가 이 제품의 신뢰성 전략")
    doc.add_page_break()

    # 12
    add_heading(doc, "슬라이드 12 — 보완점", 1)
    add_table(doc, ["영역", "보완안"], [
        ["정확도", "ARCore/ARKit SLAM, 보정 강화"],
        ["검증", "소규모 파일럿(알림 ON/OFF, 자기보고)"],
        ["UX", "위젯, 사용자 정의 임계값, 알림 채널"],
        ["프라이버시", "온디바이스 LLM 대체"],
        ["설명력", "Score 하락 요인 분해 UI"],
        ["플랫폼", "iOS IPA 빌드"],
    ])
    doc.add_page_break()

    # 13
    add_heading(doc, "슬라이드 13 — 확장 가능성 · 마무리", 1)
    add_heading(doc, "수직 확장", 2)
    add_bullets(doc, ["개인화 Score, 주간 리포트", "B2B 도서관·스터디카페 프리셋", "OS 방해금지 연동"])
    add_heading(doc, "수평 확장", 2)
    add_bullets(doc, ["센서 코어 SDK", "소음·조도 멀티모달 Score", "시설용 익명 좌석 추천"])
    add_quote(doc, "마무리: 전문 계측기를 대체하지 않고, 일상 학습에서 환경을 의식하게 만드는 스마트폰 네이티브 집중 도구")
    doc.add_paragraph("Q&A")
    doc.add_page_break()

    # appendix
    add_heading(doc, "부록 A — 발표 시간 배분 (10~11분)", 1)
    add_table(doc, ["구간", "슬라이드", "시간"], [
        ["도입", "1~2", "1분"],
        ["독창성", "3", "1분"],
        ["핵심 기능", "4~5", "2.5분"],
        ["원리", "6~8", "3분"],
        ["기술·신뢰성", "9~10", "1.5분"],
        ["한계·보완·확장", "11~13", "1.5분"],
    ])

    add_heading(doc, "부록 B — PPT AI 복붙 프롬프트", 1)
    prompt = """10분 발표. EMF Quantum Focus Android 앱. 슬라이드 13장. 미니멀 인디고 테크 스타일.

1 독창성: 스마트폰 자기장+가속도로 학습 환경 Focus Score, 집중세션/알림/AR/미니맵/구역탐색, 온디바이스, 8자보정, 민감도 프리셋.
2 핵심: 홈(게이지,집중,알림,통계), AR(카메라+EMF오버레이), 미니맵, 구역(최고 Score).
3 원리: 마그네토미터→필터/FFT→Score. 가속도→위치. AR=2D오버레이(SLAM없음). 알림=임계+히스테리시스4점.
4 스택: Flutter,Dart,provider,sensors_plus,camera,shared_preferences,선택 Gemini.
5 신뢰성: 8자보정,고정파이프라인,공식명시,히스테리시스,최소3샘플,온디바이스,한계고지. 학술은 간접(HCI,인지부하,보정). 사용자연구없음. 보장=일관성·투명성. 비보장=절대정확도·학업성과.
6 한계: 계측기아님, 1점측정, 적분오차, 가짜AR.
7 보완·확장: SLAM, 파일럿, B2B, 멀티모달."""
    p = doc.add_paragraph(prompt)
    for run in p.runs:
        run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
